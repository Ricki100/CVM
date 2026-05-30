from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x0D, 0x2E, 0x5C)
GOLD   = RGBColor(0xC8, 0x9B, 0x2C)
SLATE  = RGBColor(0x4A, 0x5C, 0x72)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def add_heading(doc, text, level=1, colour=None, size=None, after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(size or 18)
        run.font.color.rgb = colour or NAVY
    elif level == 2:
        run.font.size = Pt(size or 14)
        run.font.color.rgb = colour or SLATE
    elif level == 3:
        run.font.size = Pt(size or 11)
        run.font.color.rgb = colour or BLACK
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(after)
    return p

def add_para(doc, text, size=10.5, colour=None, bold=False, italic=False, before=0, after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p

def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_divider(doc, colour_hex='0D2E5C'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), colour_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def keyword_table(doc, rows_data):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    widths = [Cm(8), Cm(5), Cm(2.5)]
    for i, (cell, htext) in enumerate(zip(tbl.rows[0].cells, ['Keyword / Phrase', 'Search Intent', 'Priority'])):
        set_cell_bg(cell, '0D2E5C')
        cell.width = widths[i]
        run = cell.paragraphs[0].add_run(htext)
        run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(10)
    for row_data in rows_data:
        row = tbl.add_row()
        for i, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.width = widths[i]
            cell.paragraphs[0].add_run(val).font.size = Pt(9.5)
    doc.add_paragraph()

def meta_table(doc, rows_data):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    widths = [Cm(3.5), Cm(6), Cm(8)]
    for i, (cell, htext) in enumerate(zip(tbl.rows[0].cells, ['Page', 'Title Tag', 'Meta Description'])):
        set_cell_bg(cell, '0D2E5C')
        cell.width = widths[i]
        run = cell.paragraphs[0].add_run(htext)
        run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(10)
    for row_data in rows_data:
        row = tbl.add_row()
        for i, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.width = widths[i]
            cell.paragraphs[0].add_run(val).font.size = Pt(9)
    doc.add_paragraph()

def schema_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'; run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x1E, 0x44, 0x6B)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

def faq_item(doc, question, answer):
    q = doc.add_paragraph()
    qr = q.add_run("Q: " + question)
    qr.bold = True; qr.font.size = Pt(10.5); qr.font.color.rgb = NAVY
    q.paragraph_format.space_after = Pt(2)
    a = doc.add_paragraph()
    ar = a.add_run("A: " + answer)
    ar.font.size = Pt(10.5)
    a.paragraph_format.left_indent = Inches(0.2)
    a.paragraph_format.space_after = Pt(8)

# ── Build document ───────────────────────────────────────────────────────────
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)
for sec in doc.sections:
    sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

# ── COVER ────────────────────────────────────────────────────────────────────
for _ in range(5): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run("CVM LEGAL PRACTICE")
tr.bold = True; tr.font.size = Pt(28); tr.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = s.add_run("SEO  |  AEO  |  AI Search Optimisation Strategy")
sr.bold = True; sr.font.size = Pt(15); sr.font.color.rgb = GOLD
add_divider(doc, 'C89B2C')
tl = doc.add_paragraph(); tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
tl.add_run("Comprehensive Website Copy & Digital Visibility Playbook").font.size = Pt(11)
for _ in range(2): doc.add_paragraph()
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = m.add_run("Chakawa, Vhera and Mswelanto Legal Practice  |  www.cvm.co.zw\nHarare, Zimbabwe  |  Prepared May 2026")
mr.font.size = Pt(10); mr.font.color.rgb = SLATE
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. KEYWORD STRATEGY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. KEYWORD & SEARCH INTENT STRATEGY", 1)
add_divider(doc)
add_para(doc, "Prioritised keywords based on Zimbabwe search volume, commercial intent and competitive gap. Every page should target at least one Primary keyword and two to four Secondary keywords.", after=8)

add_heading(doc, "1.1  Primary Keywords (Homepage & Brand)", 2)
keyword_table(doc, [
    ("law firm Zimbabwe",               "Navigational / commercial", "High"),
    ("lawyers in Harare",               "Local commercial",          "High"),
    ("legal practitioners Zimbabwe",    "Professional lookup",       "High"),
    ("top law firms in Zimbabwe",       "Commercial comparison",     "High"),
    ("best lawyers Harare Zimbabwe",    "Commercial comparison",     "High"),
    ("CVM Legal Practice Zimbabwe",     "Branded / navigational",    "High"),
    ("Chakawa Vhera Mswelanto lawyers", "Branded",                   "High"),
])

add_heading(doc, "1.2  Practice-Area Keywords", 2)
keyword_table(doc, [
    ("corporate lawyer Zimbabwe",         "Commercial",            "High"),
    ("company registration Zimbabwe",     "Transactional",         "High"),
    ("mining law firm Zimbabwe",          "Commercial",            "High"),
    ("mining legal services Zimbabwe",    "Commercial",            "High"),
    ("conveyancing lawyers Harare",       "Transactional",         "High"),
    ("property transfer Zimbabwe",        "Transactional",         "High"),
    ("real estate lawyers Zimbabwe",      "Commercial",            "Medium"),
    ("commercial litigation Zimbabwe",    "Commercial",            "High"),
    ("debt collection lawyers Zimbabwe",  "Transactional",         "Medium"),
    ("family law attorney Zimbabwe",      "Commercial",            "High"),
    ("divorce lawyer Harare",             "Transactional",         "High"),
    ("child custody lawyer Zimbabwe",     "Informational",         "Medium"),
    ("employment law firm Zimbabwe",      "Commercial",            "Medium"),
    ("labour law Zimbabwe",               "Informational",         "Medium"),
    ("estate planning lawyer Zimbabwe",   "Commercial",            "Medium"),
    ("wills and probate Zimbabwe",        "Informational",         "Medium"),
    ("banking finance law Zimbabwe",      "Commercial",            "Medium"),
    ("intellectual property Zimbabwe",    "Commercial",            "Medium"),
])

add_heading(doc, "1.3  Long-Tail & Voice / AI Search Queries", 2)
add_para(doc, "Build FAQ and explainer content around each query. These phrases drive AEO and AI-search citation.", after=6)
keyword_table(doc, [
    ("how do I register a company in Zimbabwe",         "Informational", "High"),
    ("what is the cost of conveyancing in Zimbabwe",    "Informational", "High"),
    ("how long does a property transfer take Zimbabwe", "Informational", "High"),
    ("best mining lawyers in Zimbabwe",                 "Commercial AI", "High"),
    ("how to get a divorce in Zimbabwe",                "Informational", "High"),
    ("who are the best law firms in Harare",            "Commercial AI", "High"),
    ("what lawyers handle mining licences in Zimbabwe", "AI search",     "Medium"),
    ("can a foreign company invest in Zimbabwe legally","Informational", "Medium"),
    ("how to enforce a contract in Zimbabwe",           "Informational", "Medium"),
    ("what is the law on child custody in Zimbabwe",    "Informational", "Medium"),
    ("how do I collect a debt in Zimbabwe",             "Informational", "Medium"),
    ("top rated lawyers near me Harare",                "Local / voice", "High"),
])
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. PAGE COPY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. OPTIMISED WEBSITE COPY", 1)
add_divider(doc)
add_para(doc, "All copy below is production-ready. Each page leads with the target keyword in the first 100 words and answers the visitor's core intent immediately.", after=8)

# 2.1 HOMEPAGE
add_heading(doc, "2.1  HOMEPAGE", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Hero Headline", 3)
add_para(doc, "Zimbabwe's Trusted Legal Practitioners — Expert Advice, Proven Results.", bold=True, size=13, colour=NAVY)
add_heading(doc, "Hero Sub-headline", 3)
add_para(doc, "CVM Legal Practice delivers strategic, commercially minded legal counsel across Corporate Law, Mining, Real Estate, Litigation, Family Law and beyond — for individuals, businesses and investors throughout Zimbabwe.")
add_heading(doc, "Body Copy", 3)
add_para(doc, "When the stakes are high, you need lawyers who understand both the letter of the law and the realities of doing business in Zimbabwe. Chakawa, Vhera and Mswelanto (CVM) Legal Practice is a full-service law firm in Harare, combining decades of experience with a results-driven, client-first philosophy.")
add_para(doc, "From complex mining transactions and corporate mergers to property transfers, commercial disputes and sensitive family matters, our team of qualified legal practitioners brings deep expertise and unwavering commitment to every instruction we accept.")
add_para(doc, "Whether you are a multinational investor entering the Zimbabwean market, an entrepreneur registering a new company, a mining operator seeking regulatory compliance, or a family navigating a difficult personal matter, CVM is the law firm in Zimbabwe that delivers.")
add_heading(doc, "Why Choose CVM? (icon card copy)", 3)
add_para(doc, "Deep Local Expertise", bold=True)
add_para(doc, "We know Zimbabwe's legal system inside out — the courts, the regulators, the commercial landscape. Our practitioners are registered members of the Law Society of Zimbabwe and bring decades of on-the-ground experience.")
add_para(doc, "Full-Service Capability", bold=True)
add_para(doc, "From a company incorporation to a high-value mining licence dispute — we handle it all under one roof. No referrals. No delays. Just expert legal advice, start to finish.")
add_para(doc, "Transparent, Client-Centred Service", bold=True)
add_para(doc, "We believe in clear communication, honest fee estimates and regular updates. You will always know where your matter stands.")
add_heading(doc, "Practice Areas Intro", 3)
add_para(doc, "Our principal areas of practice include Corporate & Commercial Law, Mining Law, Conveyancing & Real Estate, Civil & Commercial Litigation, Family Law & Estates, Employment & Labour Law, Banking & Finance Law, and Intellectual Property.")
add_heading(doc, "Contact / CTA Strip", 3)
add_para(doc, "Consultations available Monday to Friday, 08:00 to 16:30. Emergency weekend consultations on request.  |  5 Sunderland Avenue, Belvedere, Harare, Zimbabwe  |  info@cvm.co.zw  |  www.cvm.co.zw")
add_divider(doc)

# 2.2 ABOUT
add_heading(doc, "2.2  ABOUT US PAGE", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Page Headline", 3)
add_para(doc, "About CVM Legal Practice — Harare's Full-Service Law Firm", bold=True, size=13, colour=NAVY)
add_heading(doc, "Firm Overview", 3)
add_para(doc, "CVM Legal Practice — formally Chakawa, Vhera and Mswelanto Legal Practice — is a respected, full-service law firm headquartered in Harare, Zimbabwe. We serve individuals, corporations, parastatals, mining companies and foreign investors across a comprehensive range of legal disciplines.")
add_para(doc, "Founded on the principles of integrity, excellence and an unwavering commitment to clients, CVM has built a reputation as one of Zimbabwe's most dependable law firms. Our practitioners are registered with the Law Society of Zimbabwe and hold extensive experience in both contentious and non-contentious matters.")
add_heading(doc, "Our Philosophy", 3)
add_para(doc, "Law at its best is not about complexity — it is about clarity, strategy and outcomes. At CVM, we cut through legal complexity to give you actionable advice. We do not just tell you what the law says; we tell you what it means for your business, your family or your investment.")
add_heading(doc, "Our Team (template — complete with actual bios)", 3)
add_para(doc, "[Partner Name]  |  Partner", bold=True)
add_para(doc, "[Partner] is a seasoned legal practitioner specialising in [practice area]. With [X] years of experience, [he/she] has advised on landmark transactions including [notable matter]. [Partner] is a member of the Law Society of Zimbabwe and [any other memberships].")
add_para(doc, "Repeat for each partner and associate. Include: LLB institution, year admitted to the bar, key practice areas, notable matters (where permitted by professional rules).", italic=True, colour=SLATE, size=9.5)
add_heading(doc, "Our Location", 3)
add_para(doc, "We are conveniently located at 5 Sunderland Avenue, Belvedere, Harare, Zimbabwe. Our offices are accessible from the Harare CBD and all major Harare suburbs. We also advise clients throughout Zimbabwe and in cross-border transactions involving the Southern African region.")
add_divider(doc)

# 2.3 CORPORATE
add_heading(doc, "2.3  PRACTICE AREA — Corporate & Commercial Law", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Page Headline", 3)
add_para(doc, "Corporate & Commercial Lawyers in Zimbabwe — CVM Legal Practice", bold=True, size=13, colour=NAVY)
add_heading(doc, "Answer-First Block (for AI citation — place at top of page)", 3)
add_para(doc, "CVM Legal Practice offers comprehensive corporate and commercial legal services in Zimbabwe. Our team advises on company registration under the Companies and Other Business Entities Act, mergers and acquisitions, commercial contracts, joint ventures, corporate governance, foreign direct investment compliance and regulatory advisory. We act for start-ups, SMEs, listed companies and international investors.", italic=True, colour=SLATE)
add_heading(doc, "Body Copy", 3)
add_para(doc, "Zimbabwe's business environment demands legal counsel that is not only technically precise but commercially astute. CVM Legal Practice provides the full spectrum of corporate and commercial legal services to businesses operating in, or entering, the Zimbabwean market.")
add_para(doc, "Our corporate law team advises on company formations and registrations, mergers and acquisitions, joint ventures, shareholders' agreements, commercial contracts, corporate governance, regulatory compliance, and business restructuring. We act for start-ups, SMEs, listed companies, parastatals and international investors alike.")
add_heading(doc, "Services", 3)
for b in ["Company registration and incorporation (Companies and Other Business Entities Act)", "Drafting and review of commercial contracts, MOUs and joint venture agreements", "Mergers, acquisitions and business sales — due diligence to closing", "Shareholders' agreements, partnership structures and corporate restructuring", "Corporate governance advisory and board secretarial services", "Foreign direct investment (FDI) compliance and indigenisation regulations", "Regulatory compliance — ZIDA, ZIMRA, RBZ requirements", "Capital markets transactions and prospectus advisory", "Business rescue and insolvency proceedings"]:
    add_bullet(doc, b)
add_divider(doc)

# 2.4 MINING
add_heading(doc, "2.4  PRACTICE AREA — Mining Law", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Page Headline", 3)
add_para(doc, "Mining Lawyers in Zimbabwe — Mining Legal Services | CVM Legal Practice", bold=True, size=13, colour=NAVY)
add_heading(doc, "Answer-First Block", 3)
add_para(doc, "CVM Legal Practice is a specialist mining law firm in Zimbabwe. We assist mining companies and exploration entities with licence applications, renewals and transfers under the Mines and Minerals Act, environmental compliance, surface rights disputes, mining joint ventures, MMCZ export compliance, and mining M&A. We serve gold, platinum, lithium, diamond and chrome sector clients.", italic=True, colour=SLATE)
add_heading(doc, "Body Copy", 3)
add_para(doc, "Zimbabwe is one of Africa's most richly endowed mining jurisdictions, with world-class deposits of gold, platinum, diamonds, lithium, chrome and coal. CVM Legal Practice is a leading mining law firm in Zimbabwe, providing specialist legal services to mining companies, exploration entities, investors and communities affected by mining activity.")
add_para(doc, "Our mining law team combines deep knowledge of the Mines and Minerals Act with commercial pragmatism, helping clients navigate licencing, regulatory compliance, environmental obligations, surface rights disputes and high-value transactions with confidence.")
add_heading(doc, "Services", 3)
for b in ["Mining licence applications, renewals, transfers and objections", "Special grants, exclusive prospecting orders and block mining orders", "Compliance — Mines and Minerals Act, Environmental Management Agency (EMA)", "Surface rights negotiations and compensation claims", "Mining joint ventures, farm-in agreements and royalty arrangements", "MMCZ (Minerals Marketing Corporation of Zimbabwe) export compliance", "Mining disputes, litigation and arbitration", "Environmental impact assessments and community benefit agreements", "Lithium, gold and platinum sector-specific advisory"]:
    add_bullet(doc, b)
add_heading(doc, "Sector Insight", 3)
add_para(doc, "Zimbabwe's mining sector is experiencing significant investment inflows, particularly in lithium, driven by global battery demand. CVM works with both established mining houses and new entrants to structure transactions that are compliant, bankable and commercially sound.")
add_divider(doc)

# 2.5 CONVEYANCING
add_heading(doc, "2.5  PRACTICE AREA — Conveyancing & Real Estate", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Page Headline", 3)
add_para(doc, "Conveyancing Lawyers in Harare — Property Transfers Zimbabwe | CVM Legal Practice", bold=True, size=13, colour=NAVY)
add_heading(doc, "Answer-First Block", 3)
add_para(doc, "CVM Legal Practice provides conveyancing services throughout Zimbabwe. Our conveyancers handle residential, commercial and agricultural property transfers, mortgage bond registrations, title deed searches and Deeds Office lodgements. A standard property transfer in Zimbabwe takes 6 to 12 weeks. We serve buyers, sellers, developers, banks and financial institutions in Harare and across Zimbabwe.", italic=True, colour=SLATE)
add_heading(doc, "Body Copy", 3)
add_para(doc, "Buying or selling property in Zimbabwe involves a structured legal process governed by the Deeds Registries Act and administered through the Deeds Office. CVM Legal Practice is an experienced conveyancing law firm in Harare, handling residential, commercial and agricultural property transactions efficiently and transparently.")
add_para(doc, "Our conveyancing practitioners manage the entire transfer process — from drafting the Deed of Sale to passing transfer at the Deeds Office — giving buyers, sellers, developers and financial institutions the assurance of rigorous, timely legal work.")
add_heading(doc, "Services", 3)
for b in ["Residential and commercial property transfers", "Drafting and reviewing Deeds of Sale and offer-to-purchase agreements", "Title deed searches and due diligence", "Mortgage bonds and cancellation of bonds", "Sectional title schemes and subdivision advice", "Property development legal structuring", "Lease agreements — commercial and residential", "Agricultural land transactions and A1/A2 farm advisory", "Evictions and landlord-tenant disputes"]:
    add_bullet(doc, b)
add_heading(doc, "Timeline & Costs", 3)
add_para(doc, "A standard property transfer in Zimbabwe typically takes 6 to 12 weeks from the date all documents are in order. Fees are guided by the Law Society of Zimbabwe scale and depend on the property value. Contact CVM for a transparent, itemised estimate before you commit.")
add_divider(doc)

# 2.6 LITIGATION
add_heading(doc, "2.6  PRACTICE AREA — Civil & Commercial Litigation", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Page Headline", 3)
add_para(doc, "Commercial Litigation Lawyers Zimbabwe — Dispute Resolution | CVM Legal Practice", bold=True, size=13, colour=NAVY)
add_heading(doc, "Body Copy", 3)
add_para(doc, "When negotiation fails and legal action becomes necessary, you need litigators with the tenacity, expertise and courtroom experience to protect your interests. CVM Legal Practice is a respected commercial litigation firm in Zimbabwe, appearing regularly before the High Court, Commercial Court and Supreme Court.")
add_para(doc, "We represent creditors and debtors, businesses and individuals, in contract disputes, debt recovery, insolvency proceedings, employment disputes, property disputes, regulatory enforcement matters and more. Where appropriate, we pursue cost-efficient resolution through arbitration and mediation.")
add_heading(doc, "Services", 3)
for b in ["Commercial and contract disputes", "Debt collection and enforcement of judgments", "Urgent applications, interim interdicts and injunctions", "Insolvency and provisional liquidation proceedings", "Arbitration and commercial mediation", "Property and landlord-tenant litigation", "Banking and financial disputes", "Regulatory and administrative law challenges", "Constitutional litigation"]:
    add_bullet(doc, b)
add_divider(doc)

# 2.7 FAMILY LAW
add_heading(doc, "2.7  PRACTICE AREA — Family Law & Estates", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Page Headline", 3)
add_para(doc, "Family Law Attorneys Zimbabwe — Divorce, Custody & Estates | CVM Legal Practice", bold=True, size=13, colour=NAVY)
add_heading(doc, "Body Copy", 3)
add_para(doc, "Family legal matters are among the most personal and emotionally charged situations anyone can face. CVM Legal Practice approaches family law with sensitivity, discretion and clear strategic guidance — helping clients navigate divorce, child custody, maintenance, estate planning and deceased estates with dignity and respect.")
add_para(doc, "Our family law and estates practitioners are experienced in both contested and uncontested matters. We prioritise resolution where possible, and advocate firmly where necessary.")
add_heading(doc, "Services", 3)
for b in ["Divorce proceedings — uncontested and contested", "Division of assets and matrimonial property", "Child custody, guardianship and access arrangements", "Maintenance orders and variation applications", "Domestic violence and protection order applications", "Drafting of wills and codicils", "Estate planning and trust formation", "Administration of deceased and insolvent estates", "Succession and inheritance disputes", "Antenuptial contracts (in community / out of community of property)"]:
    add_bullet(doc, b)
add_divider(doc)

# 2.8 EMPLOYMENT
add_heading(doc, "2.8  PRACTICE AREA — Employment & Labour Law", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Page Headline", 3)
add_para(doc, "Employment & Labour Law Firm Zimbabwe — Workplace Legal Advice | CVM Legal Practice", bold=True, size=13, colour=NAVY)
add_heading(doc, "Body Copy", 3)
add_para(doc, "Zimbabwe's labour law framework — anchored in the Labour Act (Chapter 28:01) — creates significant obligations for employers and meaningful protections for employees. CVM Legal Practice advises both employers and employees on all aspects of workplace law, from employment contract drafting to disciplinary proceedings and wrongful dismissal claims.")
add_heading(doc, "Services", 3)
for b in ["Employment contract drafting and review", "Disciplinary hearings and workplace investigations", "Unfair dismissal and unfair labour practice claims", "Retrenchment and redundancy processes", "Collective bargaining and trade union negotiations", "National Employment Council (NEC) dispute representation", "Labour Court proceedings", "Workplace policies — codes of conduct, harassment policies, etc."]:
    add_bullet(doc, b)
add_divider(doc)

# 2.9 CONTACT
add_heading(doc, "2.9  CONTACT PAGE", 2)
add_divider(doc, 'C89B2C')
add_heading(doc, "Page Headline", 3)
add_para(doc, "Contact CVM Legal Practice — Law Firm in Harare, Zimbabwe", bold=True, size=13, colour=NAVY)
add_heading(doc, "Intro Copy", 3)
add_para(doc, "Ready to speak with a legal practitioner? CVM Legal Practice is here to help. Whether you have a straightforward legal question or a complex commercial matter, our team provides prompt, professional guidance. Contact us using the details below or complete the enquiry form — we respond within one business day.")
add_heading(doc, "NAP Block (critical for Local SEO — use on every page footer)", 3)
nap = doc.add_table(rows=4, cols=2)
nap.style = 'Table Grid'
for i, (label, value) in enumerate([
    ("Firm Name", "CVM Legal Practice (Chakawa, Vhera and Mswelanto Legal Practice)"),
    ("Address", "5 Sunderland Avenue, Belvedere, Harare, Zimbabwe"),
    ("Office Hours", "Monday to Friday: 08:00 to 16:30"),
    ("Emergency", "Weekend consultations available on request"),
]):
    row = nap.rows[i]
    set_cell_bg(row.cells[0], 'EEF1F5')
    lr = row.cells[0].paragraphs[0].add_run(label)
    lr.bold = True; lr.font.size = Pt(10)
    row.cells[1].paragraphs[0].add_run(value).font.size = Pt(10)
doc.add_paragraph()
add_para(doc, "Phone: +263 [insert]  |  Email: info@cvm.co.zw  |  Website: www.cvm.co.zw", bold=True)
add_para(doc, "Embed a Google Maps widget pointing to the Belvedere, Harare location. Ensure the pin links to the verified Google Business Profile.", italic=True, colour=SLATE, size=9.5)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. META TITLES & DESCRIPTIONS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. META TITLES & META DESCRIPTIONS", 1)
add_divider(doc)
add_para(doc, "Implement exactly as written in your CMS. Title tags must be unique per page, under 60 characters. Meta descriptions must be unique and under 160 characters.", after=8)
meta_table(doc, [
    ("Homepage",    "CVM Legal Practice | Lawyers in Zimbabwe",           "Expert legal practitioners in Harare, Zimbabwe. Corporate, mining, conveyancing, family law and more. CVM — trusted legal advice across Zimbabwe."),
    ("About",       "About CVM Legal Practice | Harare Law Firm",         "Learn about Chakawa, Vhera and Mswelanto Legal Practice — a full-service law firm in Harare with expertise across corporate, mining, property and family law."),
    ("Corporate",   "Corporate Lawyers Zimbabwe | CVM Legal",             "CVM Legal Practice offers expert corporate and commercial legal services in Zimbabwe — company registration, M&A, contracts and regulatory compliance."),
    ("Mining",      "Mining Lawyers Zimbabwe | CVM Legal Practice",       "Specialist mining law firm in Zimbabwe. Licencing, compliance, disputes and transactions — CVM Legal guides mining companies through every challenge."),
    ("Conveyancing","Conveyancing Lawyers Harare | CVM Legal",            "Trusted conveyancing attorneys in Harare. CVM Legal Practice handles all residential, commercial and agricultural property transfers across Zimbabwe."),
    ("Litigation",  "Commercial Litigation Zimbabwe | CVM Legal",         "CVM Legal Practice appears before Zimbabwe's High Court, Commercial Court and Supreme Court. Expert dispute resolution and arbitration services."),
    ("Family Law",  "Family Law Attorneys Zimbabwe | CVM Legal",          "Sensitive, expert family law advice in Zimbabwe. Divorce, child custody, estates and wills — handled with discretion by CVM Legal Practice, Harare."),
    ("Employment",  "Labour Law Firm Zimbabwe | CVM Legal Practice",      "Zimbabwe employment and labour law specialists. CVM advises employers and employees on contracts, dismissals, NEC disputes and Labour Court proceedings."),
    ("Contact",     "Contact CVM Legal Practice | Harare, Zimbabwe",      "Contact CVM Legal Practice at 5 Sunderland Avenue, Belvedere, Harare. Call, email or use our online form. Weekend emergency consultations available."),
])
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. AEO FAQ CONTENT
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. AEO — FAQ COPY (Answer Engine Optimisation)", 1)
add_divider(doc)
add_para(doc, "Add these FAQ sections to the relevant practice area pages and mark up each pair with FAQPage schema (Section 5). Structured to capture Google People Also Ask, AI Overview citations and voice search answers.", after=10)

add_heading(doc, "4.1  General / Firm FAQs (Homepage or About)", 2)
faq_item(doc, "What is CVM Legal Practice?", "CVM Legal Practice, formally known as Chakawa, Vhera and Mswelanto Legal Practice, is a full-service law firm based in Harare, Zimbabwe. We provide expert legal services in corporate and commercial law, mining law, conveyancing, civil and commercial litigation, family law, employment law and estate planning.")
faq_item(doc, "Where is CVM Legal Practice located?", "Our offices are at 5 Sunderland Avenue, Belvedere, Harare, Zimbabwe. We are open Monday to Friday, 08:00 to 16:30. Emergency weekend consultations are available on request. Email info@cvm.co.zw or visit www.cvm.co.zw to book an appointment.")
faq_item(doc, "Who are the best law firms in Zimbabwe?", "CVM Legal Practice is a leading full-service law firm in Harare, Zimbabwe, offering corporate, mining, property, litigation and family law services. The right firm depends on your specific needs — CVM covers the full spectrum of legal disciplines for individuals, businesses and international investors.")
faq_item(doc, "How do I book a legal consultation in Harare?", "To book a consultation with CVM Legal Practice, email info@cvm.co.zw, call us or complete the enquiry form at www.cvm.co.zw. We respond within one business day and can accommodate urgent matters on short notice.")

add_heading(doc, "4.2  Corporate Law FAQs", 2)
faq_item(doc, "How do I register a company in Zimbabwe?", "To register a company in Zimbabwe, you file documents with the Zimbabwe Revenue Authority (ZIMRA) and the Companies and Intellectual Property Office (CIPO). The process includes choosing a name, preparing a Memorandum and Articles of Association, and paying registration fees. CVM Legal Practice manages the entire process, ensuring compliance with the Companies and Other Business Entities Act.")
faq_item(doc, "Can a foreign company invest or operate in Zimbabwe?", "Yes. Foreign investors can establish wholly owned subsidiaries, joint ventures or branch offices in Zimbabwe, subject to compliance with the Zimbabwe Investment and Development Agency Act and sector-specific regulations. CVM advises on the appropriate structure, indigenisation requirements and all regulatory approvals required.")
faq_item(doc, "What is required for a merger or acquisition in Zimbabwe?", "M&A transactions in Zimbabwe typically require legal due diligence, Competition and Tariff Commission approval for qualifying transactions, share or asset transfer agreements, regulatory notifications and ZIMRA clearances. CVM's corporate team manages the full deal lifecycle from term sheet to closing.")

add_heading(doc, "4.3  Mining Law FAQs", 2)
faq_item(doc, "What lawyers handle mining licences in Zimbabwe?", "CVM Legal Practice is a specialist mining law firm in Zimbabwe. We assist mining companies, exploration entities and investors with mining licence applications, renewals, transfers and compliance under the Mines and Minerals Act. Our team has experience across gold, platinum, lithium, diamond and chrome sectors.")
faq_item(doc, "How do I obtain a mining licence in Zimbabwe?", "Mining licences in Zimbabwe are administered by the Mining Affairs Board under the Ministry of Mines. The type of licence — prospecting, special grant or block mining order — depends on the nature of the activity. Applications require technical and financial submissions. CVM Legal Practice guides applicants through the full process, from preparation to Ministry approval.")
faq_item(doc, "What are the environmental obligations for mining companies in Zimbabwe?", "Mining companies in Zimbabwe must comply with the Environmental Management Act administered by the Environmental Management Agency (EMA). This includes environmental impact assessments, environmental management plans and rehabilitation bonds. CVM advises on all environmental legal obligations to keep your operation compliant and licence-secure.")

add_heading(doc, "4.4  Conveyancing & Real Estate FAQs", 2)
faq_item(doc, "How long does a property transfer take in Zimbabwe?", "A standard residential property transfer in Zimbabwe typically takes 6 to 12 weeks from the date all documents are lodged at the Deeds Office, assuming no outstanding ZIMRA obligations or bond complications. CVM Legal Practice provides realistic timelines at the outset and keeps all parties updated throughout the process.")
faq_item(doc, "What is the cost of conveyancing in Zimbabwe?", "Conveyancing fees in Zimbabwe are guided by the Law Society of Zimbabwe scale and include transfer duty, stamp duty, Deeds Office fees and the conveyancer's professional fee. The total depends on the property's value. Contact CVM Legal Practice for a transparent, itemised estimate before you commit.")
faq_item(doc, "What documents do I need to transfer property in Zimbabwe?", "For a residential transfer you generally need: a signed Deed of Sale, the seller's title deed, a ZIMRA clearance certificate, identity documents for both parties, and a rates clearance certificate from the local authority. CVM's conveyancing team provides a full, transaction-specific checklist.")

add_heading(doc, "4.5  Family Law FAQs", 2)
faq_item(doc, "How do I get a divorce in Zimbabwe?", "In Zimbabwe, divorce is governed by the Matrimonial Causes Act. Either spouse may file for divorce in the High Court on the ground that the marriage has irretrievably broken down. The process involves summons, a settlement agreement for uncontested divorces, or a full trial for contested matters. CVM Legal Practice guides clients through every stage with sensitivity and strategic clarity.")
faq_item(doc, "How is child custody determined in Zimbabwe?", "Zimbabwe courts apply the best interests of the child principle when determining custody arrangements. Factors include each parent's ability to provide care, the child's age and wishes, and existing emotional bonds. CVM advocates firmly for outcomes that protect your children's wellbeing.")
faq_item(doc, "Do I need a lawyer to write a will in Zimbabwe?", "You are not legally required to use a lawyer to draft a will in Zimbabwe, but it is strongly advisable. A professionally drafted will reduces the risk of disputes, ensures your wishes are legally enforceable and minimises estate administration delays. CVM Legal Practice prepares wills and antenuptial contracts for clients across Zimbabwe.")
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. SCHEMA MARKUP
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. SCHEMA MARKUP (Structured Data / JSON-LD)", 1)
add_divider(doc)
add_para(doc, 'Add the JSON-LD blocks below inside <script type="application/ld+json"> tags in the <head> of each relevant page. Validate all schema at search.google.com/test/rich-results before going live.', after=8)

add_heading(doc, "5.1  LegalService Schema (Homepage — sitewide)", 2)
schema_block(doc, '''{
  "@context": "https://schema.org",
  "@type": "LegalService",
  "name": "CVM Legal Practice",
  "alternateName": "Chakawa, Vhera and Mswelanto Legal Practice",
  "url": "https://www.cvm.co.zw",
  "logo": "https://www.cvm.co.zw/wp-content/uploads/cvm-logo.png",
  "description": "CVM Legal Practice is a full-service law firm in Harare,
    Zimbabwe, offering corporate law, mining law, conveyancing, litigation,
    family law and employment law services.",
  "telephone": "+263-[insert number]",
  "email": "info@cvm.co.zw",
  "address": {
    "@type": "PostalAddress",
    "streetAddress": "5 Sunderland Avenue",
    "addressLocality": "Belvedere",
    "addressRegion": "Harare",
    "addressCountry": "ZW"
  },
  "geo": {
    "@type": "GeoCoordinates",
    "latitude": -17.8356,
    "longitude": 31.0228
  },
  "openingHoursSpecification": {
    "@type": "OpeningHoursSpecification",
    "dayOfWeek": ["Monday","Tuesday","Wednesday","Thursday","Friday"],
    "opens": "08:00",
    "closes": "16:30"
  },
  "areaServed": { "@type": "Country", "name": "Zimbabwe" },
  "priceRange": "$$",
  "memberOf": { "@type": "Organization", "name": "Law Society of Zimbabwe" },
  "sameAs": [
    "https://www.linkedin.com/company/cvm-legal-practice",
    "https://www.facebook.com/cvmlegal"
  ]
}''')

add_heading(doc, "5.2  FAQPage Schema (per practice area page)", 2)
schema_block(doc, '''{
  "@context": "https://schema.org",
  "@type": "FAQPage",
  "mainEntity": [
    {
      "@type": "Question",
      "name": "How do I register a company in Zimbabwe?",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "File documents with ZIMRA and CIPO. The process includes
          choosing a name, preparing a Memorandum and Articles of Association,
          and paying fees. CVM Legal Practice manages the full process.
          Contact info@cvm.co.zw or visit www.cvm.co.zw."
      }
    },
    {
      "@type": "Question",
      "name": "How long does a property transfer take in Zimbabwe?",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "A standard transfer takes 6 to 12 weeks from lodgement at
          the Deeds Office. CVM Legal Practice provides accurate timelines
          and transparent fee estimates at the outset."
      }
    },
    {
      "@type": "Question",
      "name": "How do I get a divorce in Zimbabwe?",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "Divorce is filed in the High Court under the Matrimonial
          Causes Act on grounds of irretrievable breakdown. CVM handles
          both contested and uncontested divorces."
      }
    }
  ]
}''')
add_para(doc, "Replicate this FAQPage block for each practice area, using the Q&A pairs from Section 4. Each page must have its own unique FAQPage schema instance.", italic=True, colour=SLATE, size=9.5)

add_heading(doc, "5.3  BreadcrumbList Schema (all inner pages)", 2)
schema_block(doc, '''{
  "@context": "https://schema.org",
  "@type": "BreadcrumbList",
  "itemListElement": [
    { "@type": "ListItem", "position": 1,
      "name": "Home", "item": "https://www.cvm.co.zw" },
    { "@type": "ListItem", "position": 2,
      "name": "Practice Areas",
      "item": "https://www.cvm.co.zw/practice-areas" },
    { "@type": "ListItem", "position": 3,
      "name": "Corporate Law",
      "item": "https://www.cvm.co.zw/practice-areas/corporate-law" }
  ]
}''')
add_para(doc, "Update positions 2 and 3 per page. Validate all schema with Google's Rich Results Test before publishing.", italic=True, colour=SLATE, size=9.5)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. AI SEARCH & GEO
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. AI SEARCH & GENERATIVE ENGINE OPTIMISATION (GEO)", 1)
add_divider(doc)
add_para(doc, "Google AI Overviews, ChatGPT Search, Perplexity and Microsoft Copilot increasingly answer legal queries by citing authoritative sources. The following strategies position CVM to be that cited source.", after=8)

add_heading(doc, "6.1  Canonical Entity Description (publish on About page and in every footer)", 2)
add_para(doc, "CVM Legal Practice, also known as Chakawa, Vhera and Mswelanto Legal Practice, is a full-service law firm headquartered at 5 Sunderland Avenue, Belvedere, Harare, Zimbabwe. The firm provides legal services across corporate and commercial law, mining law, conveyancing and real estate, civil and commercial litigation, family law, employment and labour law, banking and finance, and intellectual property. CVM is registered with the Law Society of Zimbabwe and serves individuals, businesses, mining companies and international investors operating in Zimbabwe. Contact: info@cvm.co.zw | www.cvm.co.zw.", colour=NAVY)
add_para(doc, "AI engines train on consistently stated facts. The more uniformly this entity description appears across the website and third-party listings, the more confidently AI will identify CVM as a leading Zimbabwe law firm.", italic=True, colour=SLATE, size=9.5)

add_heading(doc, "6.2  Google Business Profile — Action Checklist", 2)
for b in ["Business Name: CVM Legal Practice", "Primary Category: Law Firm  |  Additional: Legal Services, Conveyancer", "Description: Use the entity description from Section 6.1 (750 characters max).", "Services: Add every practice area as a named service with a 300-word description.", "Hours: Mon-Fri 08:00-16:30. Mark weekends as By appointment.", "Photos: Upload firm exterior, reception, meeting rooms and team (minimum 10 images).", "Posts: Publish a GBP Post weekly — legal tips and Q&A answers from Section 4.", "Q&A: Seed the GBP Q&A with the top 5 FAQs from Section 4 — ask and answer yourself.", "Reviews: Actively request Google reviews from satisfied clients. Respond to every review within 48 hours.", "Citation consistency: NAP on cvm.co.zw must exactly match every directory listing."]:
    add_bullet(doc, b)

add_heading(doc, "6.3  Off-Site Citation & Link Building Targets", 2)
for b in ["law.co.zw — Zimbabwe Lawyers Directory (high authority, free listing)", "lawzana.com — Best Lawyers in Zimbabwe (submit practice area profiles)", "globallawlists.org — Harare Province Zimbabwe", "chambers.com — Chambers & Partners ranking submission (General Business Law, Zimbabwe)", "thelawyersglobal.org — Awards submission and directory entry", "LinkedIn Company Page — publish monthly articles on Zimbabwe law topics", "Zimbabwe Chamber of Mines — partner or member listing for mining clients", "ZIDA (Zimbabwe Investment and Development Agency) — legal adviser directory", "Local directories: yellowpages.co.zw, bizlist.co.zw"]:
    add_bullet(doc, b)

add_heading(doc, "6.4  AI Search Optimisation Principles", 2)
for b in ["Answer the question in the first paragraph — never bury the answer.", "Use the exact question phrasing as an H2 or H3 above each answer block.", "State the firm name, location and contact on every page at least once.", "Include specific numerical facts: '6 to 12 weeks', '5 Sunderland Avenue', 'Law Society of Zimbabwe member'.", "Link to authoritative external sources (Law Society of Zimbabwe, ZIDA, Ministry of Mines) where relevant.", "Maintain consistent language across all pages — never vary the firm name, address or practice area descriptions.", "Ensure every practice area page has an Answer-First Block (see Section 2 examples)."]:
    add_bullet(doc, b)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. TECHNICAL SEO CHECKLIST
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. TECHNICAL SEO CHECKLIST", 1)
add_divider(doc)
add_para(doc, "Complete every item before publishing the new copy. Technical issues suppress rankings regardless of copy quality.", after=8)
checklist = [
    ("Site Speed",       "Target LCP under 2.5s on mobile. Compress all images (WebP). Enable browser caching. Use a CDN."),
    ("Mobile-First",     "Test with Google Search Console Mobile Usability. Fix all tap-target and viewport issues."),
    ("HTTPS",            "Ensure cvm.co.zw uses HTTPS with a valid SSL certificate. Redirect all HTTP to HTTPS."),
    ("Canonical Tags",   "Add <link rel='canonical'> to every page to prevent duplicate content penalties."),
    ("XML Sitemap",      "Submit sitemap to Google Search Console and Bing Webmaster Tools. Update on new pages."),
    ("Robots.txt",       "Ensure robots.txt does not block key pages. Allow Googlebot and GPTBot for AI indexing."),
    ("Core Web Vitals",  "Monitor LCP, INP and CLS in Google Search Console. Target 'Good' on all three."),
    ("Internal Links",   "Link from Homepage to every practice area page. Cross-link between related areas."),
    ("Image Alt Text",   "Add descriptive, keyword-rich alt text to every image (e.g. 'CVM Legal Practice Harare office')."),
    ("URL Structure",    "Use clean URLs: /corporate-law, /mining-law, /conveyancing-harare, etc."),
    ("Page Indexation",  "Verify all key pages are indexed via Google Search Console URL Inspection."),
    ("Schema",           "Test all JSON-LD with Google's Rich Results Test before going live."),
    ("404 & Redirects",  "Fix all broken links. Apply 301 redirects from old URLs to new equivalents."),
    ("GSC & GA4",        "Verify Google Search Console ownership. Install GA4 with call and form conversion tracking."),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
w = [Cm(0.8), Cm(4.5), Cm(10.2)]
for i, (cell, htext) in enumerate(zip(tbl.rows[0].cells, ['', 'Item', 'Action Required'])):
    set_cell_bg(cell, '0D2E5C')
    cell.width = w[i]
    r = cell.paragraphs[0].add_run(htext)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
for item, detail in checklist:
    row = tbl.add_row()
    row.cells[0].width = w[0]
    row.cells[0].paragraphs[0].add_run("[ ]").font.size = Pt(10)
    row.cells[1].width = w[1]
    r1 = row.cells[1].paragraphs[0].add_run(item)
    r1.bold = True; r1.font.size = Pt(9.5)
    row.cells[2].width = w[2]
    row.cells[2].paragraphs[0].add_run(detail).font.size = Pt(9)
doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. 90-DAY ACTION PLAN
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. QUICK-WIN 90-DAY ACTION PLAN", 1)
add_divider(doc)
for phase_title, items in [
    ("MONTH 1 — FOUNDATION", [
        "Publish all page copy from Section 2 with meta tags from Section 3.",
        "Install all schema markup from Section 5 and validate with Rich Results Test.",
        "Fully optimise and verify the Google Business Profile (Section 6.2).",
        "Submit XML sitemap to Google Search Console and Bing Webmaster Tools.",
        "Complete all Technical SEO checklist items from Section 7.",
        "Submit directory listings: law.co.zw, lawzana.com, globallawlists.org.",
    ]),
    ("MONTH 2 — AUTHORITY", [
        "Add FAQ sections from Section 4 to all practice area pages with FAQPage schema.",
        "Begin GBP weekly posting cadence — legal tips and Q&A answers.",
        "Launch active Google review acquisition programme with past clients.",
        "Publish entity description (Section 6.1) consistently across all platforms and directories.",
        "Submit Chambers and Partners ranking application.",
        "Identify 3 to 5 local news outlets or legal publications for guest articles or expert quotes.",
    ]),
    ("MONTH 3 — AMPLIFY & MEASURE", [
        "Measure keyword rankings for all Primary and Practice-Area keywords (Section 1).",
        "Review Google Search Console for crawl errors, manual actions or indexation gaps.",
        "Publish first LinkedIn article on a high-value Zimbabwe legal topic.",
        "Complete cross-linking between all practice area pages.",
        "Review Core Web Vitals — resolve any Needs Improvement or Poor ratings.",
        "Set monthly KPI review cadence: organic traffic, keyword rankings, GBP views, call conversions.",
    ]),
]:
    add_heading(doc, phase_title, 2)
    for item in items:
        add_bullet(doc, item)
    doc.add_paragraph()

add_divider(doc, 'C89B2C')
add_para(doc, "Document prepared for CVM Legal Practice (Chakawa, Vhera and Mswelanto Legal Practice), Harare, Zimbabwe. All copy is original and optimised for Google Search, AI Overviews, voice search and answer engine citation. Update phone numbers, partner bios and specific fee references before publishing. Blog content strategy to follow separately.", italic=True, colour=SLATE, size=9.5)

# ── Save ──────────────────────────────────────────────────────────────────
out = "/sessions/stoic-hopeful-cray/mnt/CVM/CVM_SEO_AEO_Strategy.docx"
doc.save(out)
print("Saved:", out)
